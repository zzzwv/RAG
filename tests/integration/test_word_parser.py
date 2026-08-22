from io import BytesIO

from docx import Document

from rag_app.parsing.word_parser import WordParser


def test_docx_paragraphs_and_tables_are_extracted():
    document = Document()
    document.add_paragraph("员工手册")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "年假"
    table.cell(0, 1).text = "五天"
    payload = BytesIO()
    document.save(payload)
    parsed = WordParser().parse("handbook.docx", payload.getvalue())
    assert "员工手册" in parsed.text
    assert "年假 | 五天" in parsed.text
